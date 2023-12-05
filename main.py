import pandas as pd
import customtkinter
from docx import Document
from docx.shared import Pt
from docx.shared import Cm
from docx.enum.text import WD_LINE_SPACING
from docx.enum.text import WD_ALIGN_PARAGRAPH
from tkinter import *
from tkinter import filedialog
from tkinter import messagebox
from PIL import Image, ImageTk

# source_path = r"C:\Users\Nutzer\Desktop\Klausuren_Test_1.xlsx"

# -------------------Bausteine der GUI------------------------

# Erzeugt GUI
root = customtkinter.CTk()

# Design der GUI
customtkinter.set_appearance_mode("dark")  # Modes: system (default), light, dark
customtkinter.set_default_color_theme(
    "dark-blue"
)  # Themes: blue (default), dark-blue, green

# Name für GUI
root.title("")

# Icon für GUI
root.iconbitmap(
    r"C:\Users\Nutzer\PythonProjekte\Verschiedenes\TestScoreDOCX\images\excel_icon.ico"
)

# Größe des GUI-Fensters
root.geometry("300x250")
root.minsize(300, 250)
root.maxsize(300, 250)

# Icons für Buttons
add_list_image = customtkinter.CTkImage(
    Image.open("images/add-list.png").resize((30, 30), Image.Resampling.LANCZOS)
)
add_folder_image = customtkinter.CTkImage(
    Image.open("images/add-folder.png").resize((30, 30), Image.Resampling.LANCZOS)
)
process_file_image = customtkinter.CTkImage(
    Image.open("images/process-file.png").resize((30, 30), Image.Resampling.LANCZOS)
)

# Fenster für Statusbalken


# def progress_window():
#     progress = Toplevel(root)
#     progress.title("")
#     progress.geometry("250x100")
#     progress.minsize(250, 100)
#     progress.maxsize(250, 100)
#     progressbar = customtkinter.CTkProgressBar(master=progress)
#     progressbar.configure(width=150, determinate_speed=2)
#     progressbar.set(0)
#     progressbar.place(relx=0.5, rely=0.5, anchor=customtkinter.CENTER)
#     progressbar.start()


# Globale Variable, in der der Pfad zur Excel-Datei gespeichert wird
source_path = ""
dest_path = ""

# Flag, die anzeigt, ob eine Quell-Datei gewählt wurde
file_chosen = False

# Flag, die anzeigt, ob ein Zielordner gewählt wurde
dir_chosen = False


# -----------------Funktionen der GUI---------------------


def choose_file():
    """
    Funktion, die aufgerufen wird, wenn man den Button zum Auswählen der Excel-Mappe drückt.
    """

    # Pfad für die Exceldatei vom Benutzer angeben lassen und in einer Variablen speichern
    root.filename = filedialog.askopenfilename(
        initialdir="C://Users/Nutzer/Desktop",
        title="Excel-Datei wählen",
        filetypes=[("Excel Dateien", "*.xlsx")],
    )
    global source_path
    global show_source_path
    source_path = root.filename

    # Sorgt dafür, dass wirklich eine Datei gewählt wurde
    if source_path != "":
        global file_chosen
        file_chosen = True


def choose_dest():
    """
    Funktion, die aufgerufen wird, wenn man den Button zum Speichern der fertigen DOCX drückt.
    """

    # Auswahl eines Zielordners nur erlauben, wenn auch eine Quelldatei gewählt wurde
    if file_chosen:
        # Pfad für den Zielordner vom Benutzer angeben lassen und in einer Variablen speichern
        root.filename = filedialog.askdirectory(
            initialdir="C://Users/Nutzer/Desktop", title="Zielordner wählen"
        )
        global dest_path
        dest_path = root.filename

        if dest_path != "":
            global dir_chosen
            dir_chosen = True

    else:
        messagebox.showinfo(title="Datei", message="Es wurde noch keine Datei gewählt.")


def create_DOCX():
    """
    Diese Funktion wird aufgerufen, wenn der Button zum Erstellen der DOCX gedrückt wird.
    Hier passiert dasselbe wie in dem Programm, das ohne GUI läuft.
    """

    global file_chosen
    global dir_chosen

    if file_chosen and dir_chosen:
        # ---------------------Auslesen der Excel-Tabelle-----------------------

        # Erzeugen eines Data-Frames.
        df = pd.read_excel(source_path)
        # df = df.dropna(subset="Aufgabe")

        # Erste Datenreihe (also die zweite Zeile des Blattes). Gibt eine Serie zurück.
        # Dabei ist der ersten Zeile des Blattes (quasi die Überschrift jeder Spalte)
        # immer der Wert der Datenreihe zugeordnet, die hier in der eckigen Klammer steht.
        row_one = df.iloc[0]

        row_two = df.iloc[1]  # Zweite Zeile der Excel-Tabelle

        row_four = df.iloc[3]  # Dritte Zeile der Excel-Tabelle

        # Die maximal erreichbare Punktzahl.
        # Ist so zu lesen: Gib mir den Wert aus der ersten Datenreihe, der in der Spalte
        # mit der Überschrift "Maximal erreichbare Gesamtpunktzahl" steht.
        points_max = row_one["Maximal erreichbare Gesamtpunktzahl"]

        if (
            str(points_max)[-1] == "0"
        ):  # Falls .0, dann soll die Null nicht ausgegeben werden.
            points_max = int(points_max)

        # Erstellung eines Dictionary, in dem jeder Aufgabe die zu erreichende Punktzahl zugeordnet wird:
        df_points = df[["Name Aufgabe", "Maximal erreichbare Punktzahl pro Aufgabe"]]

        points_per_excercise_max = {}

        for pos, data in df_points.iterrows():
            if (
                pd.isna(data[0]) or data[0] == 0
            ):  # Falls keine Zahl eingetragen ist oder die Zahl eine Null ist, überspringen.
                pass
            # Überprüfen, ob die zu erreichende Punktzahl eine Dezimalstelle hat.
            elif str(data[1])[-1] == "0":
                # Wenn z.B. 9.0 zu erreichen, dann soll nur 9 in der Tabelle stehen.
                points_per_excercise_max[data[0]] = int(
                    data[1]
                )  # Eintrag ins Dictionary.
            else:
                points_per_excercise_max[data[0]] = data[1]  # Eintrag ins Dictionary.

        # Liste der Aufgaben, mit den Namen, wie sie auch in der Arbeit vergeben wurden:
        exercises = []
        for key in points_per_excercise_max:
            key = str(key)
            if len(key) > 1:
                if key[-2] == ".":
                    key = key[:-1]
            exercises.append(str(key))

        # Notenspiegel (die Noten stehen in den Spalten 31-37 der Excel-Tabelle)
        overview_of_grades = [int(i) for i in row_two[31:37]]
        average_score = round(row_four[31], 1)

        # Liste erstellen, in der nur die Rohpunkte in Prozent enthalten sind
        df_percent = df[
            ["Notenschlüssel"]
        ]  # Neues df, in dem nur die Rohpunkte in Prozent relevant sind

        grading_scale_percent = []

        for pos, data in df_percent.iterrows():
            if pos == 0 or pos > 16:
                pass
            else:
                grading_scale_percent.append(data[0])

        # Liste erstellen, in der nur die zu den Rohpunkten passenden MSS-Noten enthalten sind
        df_mss = df[
            ["Unnamed: 39"]
        ]  # Neues df, in dem nur die MSS-Punkte relevant sind

        grading_scale_mss = []

        for pos, data in df_mss.iterrows():
            if pos == 0 or pos > 16:
                pass
            else:
                grading_scale_mss.append(data[0])

        # Liste erstellen, in der nur die zu den Rohpunkten passenden Noten in Worten enthalten sind
        df_grades_in_words = df[
            ["Unnamed: 40"]
        ]  # Neues df, in dem nur die Noten in Worten relevant sind

        grading_scale_words = []

        for pos, data in df_grades_in_words.iterrows():
            if pos == 0 or pos > 16:
                pass
            else:
                grading_scale_words.append(data[0])

        document = Document()

        # --------------------Funktion zur Erstellung der DOCX------------------------

        def create_student(
            name,
            mss_note,
            note,
            anzahl_aufgaben,
            punkte_pro_aufgabe,
            punkte_aufgaben_erreicht,
            punkte_gesamt,
            erreicht_gesamt,
            prozent,
        ):
            # --------------------------Seitenränder der DOCX--------------------------

            sections = document.sections
            for section in sections:
                section.top_margin = Cm(1)
                section.bottom_margin = Cm(1)

            # -----------------------------Kopf der DOCX------------------------------

            title = document.add_paragraph()
            title.add_run(name).font.size = Pt(26)
            title.paragraph_format.line_spacing = 1
            title.paragraph_format.space_after = Pt(1)
            line = document.add_paragraph(
                "_________________________________________________________________________________________________________"
            )
            line.paragraph_format.space_before = Pt(1)
            paragraph = document.add_paragraph()
            paragraph.add_run("MSS-Punkte: ").font.size = Pt(16)
            grade = paragraph.add_run(str(mss_note))
            grade.bold = True
            grade.font.size = Pt(16)
            paragraph.add_run("             ")
            paragraph.add_run("Note: ").font.size = Pt(16)
            grade = paragraph.add_run(note)
            grade.bold = True
            grade.font.size = Pt(16)
            document.add_paragraph().add_run("Punkteverteilung:").font.size = Pt(14)

            # ------------------Ergebnisse der einzelnen Aufgaben--------------------

            # Um die Formatierung in der Tabelle beizubehalten, darf diese nicht zu viele Spalten haben.
            # Als Grenze wird eine Länge von 11 Spalten verwendet. Ist die Tabelle länger, wird sie aufgeteilt.

            if len(exercises) <= 11:
                table = document.add_table(rows=3, cols=anzahl_aufgaben + 2)
                table.style = "Light Shading"
                # document.add_paragraph()

                for row in table.rows:
                    for cell in row.cells:
                        if cell == table.rows[0].cells[0]:
                            pass
                        else:
                            cell.text = ""
                table.rows[1].cells[0].text = "Punkte maximal"
                table.rows[2].cells[0].text = "Punkte erreicht"

                for i in range(len(table.columns)):
                    if i == 0:
                        pass
                    elif i == (len(table.columns) - 1):
                        table.rows[0].cells[i].text = "Gesamt"
                    else:
                        table.rows[0].cells[i].text = str(exercises[i - 1])

                table_column_index = 1
                for item in punkte_pro_aufgabe:
                    table.rows[1].cells[table_column_index].text = str(
                        punkte_pro_aufgabe[item]
                    )
                    table_column_index += 1

                table_column_index_2 = 1
                for item in punkte_aufgaben_erreicht:
                    table.rows[2].cells[table_column_index_2].text = str(
                        punkte_aufgaben_erreicht[item]
                    )
                    table_column_index_2 += 1

                if str(erreicht_gesamt)[-1] == "0":
                    erreicht_gesamt = int(erreicht_gesamt)

                table.rows[1].cells[(len(table.columns) - 1)].text = str(punkte_gesamt)
                table.rows[2].cells[(len(table.columns) - 1)].text = str(
                    erreicht_gesamt
                )

            else:
                table_1 = document.add_table(rows=3, cols=12)
                table_1.style = "Light Shading"
                document.add_paragraph()
                table_2 = document.add_table(rows=3, cols=anzahl_aufgaben - 9)
                table_2.style = "Light Shading"
                document.add_paragraph()

                # Füllen der ersten Tabelle.
                for row in table_1.rows:
                    for cell in row.cells:
                        if cell == table_1.rows[0].cells[0]:
                            pass
                        else:
                            cell.text = ""

                table_1.rows[1].cells[0].text = "Punkte maximal"
                table_1.rows[2].cells[0].text = "Punkte erreicht"

                for i in range(len(table_1.columns)):
                    if i == 0:
                        pass
                    else:
                        table_1.rows[0].cells[i].text = str(exercises[i - 1])

                table_column_index = 1
                for index, item in enumerate(punkte_pro_aufgabe):
                    if index <= 10:
                        table_1.rows[1].cells[table_column_index].text = str(
                            punkte_pro_aufgabe[item]
                        )
                        table_column_index += 1

                table_column_index_2 = 1
                for index, item in enumerate(punkte_aufgaben_erreicht):
                    if index <= 10:
                        table_1.rows[2].cells[table_column_index_2].text = str(
                            punkte_aufgaben_erreicht[item]
                        )
                        table_column_index_2 += 1

                # Füllen der zweiten Tabelle.
                for row in table_2.rows:
                    for cell in row.cells:
                        if cell == table_2.rows[0].cells[0]:
                            pass
                        else:
                            cell.text = ""
                table_2.rows[1].cells[0].text = "Punkte maximal"
                table_2.rows[2].cells[0].text = "Punkte erreicht"

                for i in range(len(table_2.columns)):
                    if i == 0:
                        pass
                    elif i == (len(table_2.columns) - 1):
                        table_2.rows[0].cells[i].text = "Gesamt"
                    else:
                        table_2.rows[0].cells[i].text = str(exercises[i + 10])

                table_column_index = 1
                for index, item in enumerate(punkte_pro_aufgabe):
                    if index > 10:
                        table_2.rows[1].cells[table_column_index].text = str(
                            punkte_pro_aufgabe[item]
                        )
                        table_column_index += 1

                table_column_index_2 = 1
                for index, item in enumerate(punkte_aufgaben_erreicht):
                    if index > 10:
                        table_2.rows[2].cells[table_column_index_2].text = str(
                            punkte_aufgaben_erreicht[item]
                        )
                        table_column_index_2 += 1

                table_2.rows[1].cells[(len(table_2.columns) - 1)].text = str(
                    punkte_gesamt
                )

                if str(erreicht_gesamt)[-1] == "0":
                    erreicht_gesamt = int(erreicht_gesamt)

                table_2.rows[2].cells[(len(table_2.columns) - 1)].text = str(
                    erreicht_gesamt
                )

            percentage = document.add_paragraph().add_run(
                "Die von dir erreichte Gesamtpunktzahl entspricht "
                + str(prozent)
                + " % der maximal erreichbaren Punkte."
            )
            percentage.font.size = Pt(14)

            # ---------------------------Notenspiegel----------------------------

            # document.add_paragraph()
            document.add_paragraph().add_run("Notenspiegel:").font.size = Pt(14)
            table_overview = document.add_table(rows=3, cols=7)
            table_overview.style = "Light Shading"
            table_overview.rows[0].cells[0].text = "Note"
            table_overview.rows[0].cells[1].text = "1"
            table_overview.rows[0].cells[2].text = "2"
            table_overview.rows[0].cells[3].text = "3"
            table_overview.rows[0].cells[4].text = "4"
            table_overview.rows[0].cells[5].text = "5"
            table_overview.rows[0].cells[6].text = "6"
            table_overview.rows[1].cells[0].text = "Anzahl"
            table_overview.rows[2].cells[0].text = "Durchschnitt"
            table_overview.rows[2].cells[1].text = str(average_score)

            for i in range(len(table_overview.columns)):
                if i == 0:
                    pass
                else:
                    table_overview.rows[1].cells[i].text = str(
                        overview_of_grades[i - 1]
                    )

            # --------------------------Notenschlüssel--------------------------

            document.add_paragraph()
            document.add_paragraph().add_run(
                "Angewendeter Notenschlüssel:"
            ).font.size = Pt(14)
            table_grading_scale = document.add_table(rows=17, cols=3)
            table_grading_scale.style = "Light Shading"
            table_grading_scale.alignment = WD_ALIGN_PARAGRAPH.CENTER
            table_grading_scale.rows[0].cells[0].text = "Rohpunkte in %"
            table_grading_scale.rows[0].cells[1].text = "MSS-Punkte"
            table_grading_scale.rows[0].cells[2].text = "Note in Worten"

            for i, element in enumerate(grading_scale_percent):
                table_grading_scale.rows[i + 1].cells[0].text = str(element)

            for i, element in enumerate(grading_scale_mss):
                table_grading_scale.rows[i + 1].cells[1].text = str(element)

            for i, element in enumerate(grading_scale_words):
                table_grading_scale.rows[i + 1].cells[2].text = str(element)

        # -------------Wiederholtes Aufrufen von create_student()-------------

        row_index = 2  # Laufindex, um Daten zu jedem Schüler zu sammeln (Start bei 2, da bei 0 und 1 die die Aufgaben stehen)

        for row in range(len(df) - 2):
            # Die nächste Zeile fängt zwei Fälle ab:
            # 1. Am Ende der Schülerliste folgen nur noch Nullen. Hier soll nichts geschehen.
            # 2. Wenn ein Schüler nicht mitgeschrieben hat, steht in jeder Zelle seiner Zeile NaN.
            # Eine solche Zeile muss übersprungen werden, damit es keinen Fehler gibt.
            if df.iloc[row_index, 0] == 0 or pd.isna(df.iloc[row_index, 1]):
                pass
            else:
                # Dictionary, in dem einem Schüler zugeordnet wird, wie viele Punkte er pro Aufgabe erhalten hat:
                name = df.iloc[row_index, 0]  # Name des Schülers der aktuellen Zeile.
                # Nur der Teil der aktuellen Zeile, in dem die Ergebnisse der einzelnen Aufgaben stehen
                student = df.iloc[row_index, 1 : len(exercises) + 1]
                score_per_exercise = {}

                for index, item in enumerate(student):
                    grade = str(item)
                    # Falls 0 Punkte erreicht wurden, wird dies eingetragen
                    if grade == "0":
                        score_per_exercise[exercises[index]] = grade
                    # Falls man z.B. 5.0 erreicht hat, soll nur 5 eingetragen werden
                    elif grade[-1] == "0":
                        score_per_exercise[exercises[index]] = grade[:-2]
                    # In allen anderen Fällen wird der tatsächliche Wert übernommen
                    else:
                        score_per_exercise[exercises[index]] = grade

                # Dictionary, in dem einem Schüler die erreichte Gesamtpunktzahl, die Prozent der zu erreichenden Gesamtpunktzahl,
                # die MSS-Punkte, und die Note in Worten zugeordnet werden
                student_2 = df.loc[
                    row_index,
                    ["Punkte gesamt", "Prozent gesamt", "MSS-Punkte", "Note in Worten"],
                ]
                labels = [
                    "Punkte gesamt",
                    "Prozent gesamt",
                    "MSS-Punkte",
                    "Note in Worten",
                ]
                index = 0
                results = {}
                for item in student_2:
                    results[labels[index]] = item
                    index += 1

                create_student(
                    name,
                    int(results["MSS-Punkte"]),
                    results["Note in Worten"],
                    len(exercises),
                    points_per_excercise_max,
                    score_per_exercise,
                    points_max,
                    results["Punkte gesamt"],
                    int(results["Prozent gesamt"]),
                )
                document.add_page_break()  # Sorgt dafür, dass mit jedem Schüler eine neue Seite angefangen wird.
            row_index += 1

        # document.save(r"C:\Users\Nutzer\Desktop\Klausurergebnisse.docx")
        document.save(
            dest_path + "//Klausurergebnisse.docx"
        )  # Erstellt am Ende die eigentliche DOCX.

    else:
        message = messagebox.showinfo(
            "Datei und Ziel", "Zuerst Datei und und Zielordner wählen!"
        )


# Button zum Starten der Dateiauswahl
button_file = customtkinter.CTkButton(
    master=root,
    text="Datei wählen",
    image=add_list_image,
    command=choose_file,
    width=190,
    height=40,
    compound=RIGHT,
)

# Button zum Starten der Zielordnerauswahl
button_dir = customtkinter.CTkButton(
    master=root,
    text="Zielordner wählen",
    image=add_folder_image,
    command=choose_dest,
    width=190,
    height=40,
    compound=RIGHT,
)

# Button, um die Erstellung der Datein zu starten
button_start = customtkinter.CTkButton(
    master=root,
    text="DOCX erzeugen",
    image=process_file_image,
    command=create_DOCX,
    width=190,
    height=40,
    compound=RIGHT,
)


button_file.place(relx=0.5, rely=0.3, anchor=customtkinter.CENTER)

button_dir.place(relx=0.5, rely=0.5, anchor=customtkinter.CENTER)

button_start.place(relx=0.5, rely=0.7, anchor=customtkinter.CENTER)

# progressbar.place(relx=0.5, rely=0.9, anchor=customtkinter.CENTER)


root.mainloop()
